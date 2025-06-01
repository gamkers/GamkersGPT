
import streamlit as st
import base64
import os
import requests
import json
import time
import re
from datetime import datetime, timedelta
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
import requests
import xml.etree.ElementTree as ET
from st_on_hover_tabs import on_hover_tabs
import streamlit.components.v1 as components
from typing import List, Dict
from langchain.chains import LLMChain, ConversationChain
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
import json
from docx import Document
import subprocess
import tempfile
from streamlit_lottie import st_lottie
from youtubesearchpython import VideosSearch
from googlesearch import search



# Rate limiting configuration
MAX_API_CALLS = 10
RESET_PERIOD_HOURS = 24  # Reset limit every 24 hours

def get_user_id():
    """Generate a unique user ID based on session"""
    if 'user_id' not in st.session_state:
        # Create a unique ID based on session info and timestamp
        session_info = str(st.session_state) + str(time.time())
        st.session_state.user_id = hashlib.md5(session_info.encode()).hexdigest()[:12]
    return st.session_state.user_id

def initialize_rate_limiter():
    """Initialize rate limiting session state variables"""
    if 'api_call_count' not in st.session_state:
        st.session_state.api_call_count = 0
    
    if 'first_api_call_time' not in st.session_state:
        st.session_state.first_api_call_time = None
    
    if 'rate_limit_exceeded' not in st.session_state:
        st.session_state.rate_limit_exceeded = False

def check_rate_limit_reset():
    """Check if rate limit should be reset based on time"""
    if st.session_state.first_api_call_time:
        time_since_first_call = datetime.now() - st.session_state.first_api_call_time
        if time_since_first_call > timedelta(hours=RESET_PERIOD_HOURS):
            # Reset the rate limit
            st.session_state.api_call_count = 0
            st.session_state.first_api_call_time = None
            st.session_state.rate_limit_exceeded = False
            st.success("🔄 Your API usage limit has been reset!")

def can_make_api_call():
    """Check if user can make an API call"""
    check_rate_limit_reset()
    return st.session_state.api_call_count < MAX_API_CALLS

def increment_api_call():
    """Increment API call counter"""
    if st.session_state.first_api_call_time is None:
        st.session_state.first_api_call_time = datetime.now()
    
    st.session_state.api_call_count += 1
    
    if st.session_state.api_call_count >= MAX_API_CALLS:
        st.session_state.rate_limit_exceeded = True

def display_rate_limit_info():
    """Display current rate limit status"""
    remaining_calls = MAX_API_CALLS - st.session_state.api_call_count
    
    if st.session_state.first_api_call_time:
        reset_time = st.session_state.first_api_call_time + timedelta(hours=RESET_PERIOD_HOURS)
        time_until_reset = reset_time - datetime.now()
        
        if time_until_reset.total_seconds() > 0:
            hours, remainder = divmod(int(time_until_reset.total_seconds()), 3600)
            minutes, _ = divmod(remainder, 60)
            reset_info = f"Resets in: {hours}h {minutes}m"
        else:
            reset_info = "Ready to reset"
    else:
        reset_info = "No calls made yet"
    
    # Create a nice info box
    if remaining_calls > 5:
        status_color = "🟢"
    elif remaining_calls > 2:
        status_color = "🟡"
    else:
        status_color = "🔴"
    
    st.sidebar.markdown(f"""
    <div style="
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #10b981;
        margin-bottom: 1rem;
    ">
        <h4 style="margin: 0 0 0.5rem 0; color: #e2e8f0;">
            {status_color} API Usage Limit
        </h4>
        <p style="margin: 0; color: #94a3b8;">
            <strong>Remaining:</strong> {remaining_calls}/{MAX_API_CALLS} calls<br>
            <strong>User ID:</strong> {get_user_id()}<br>
            <strong>Status:</strong> {reset_info}
        </p>
    </div>
    """, unsafe_allow_html=True)

def rate_limit_warning():
    """Display rate limit exceeded warning"""
    if st.session_state.rate_limit_exceeded:
        reset_time = st.session_state.first_api_call_time + timedelta(hours=RESET_PERIOD_HOURS)
        time_until_reset = reset_time - datetime.now()
        
        if time_until_reset.total_seconds() > 0:
            hours, remainder = divmod(int(time_until_reset.total_seconds()), 3600)
            minutes, _ = divmod(remainder, 60)
            
            st.error(f"""
            🚫 **API Rate Limit Exceeded**
            
            You have reached the maximum of {MAX_API_CALLS} API calls per {RESET_PERIOD_HOURS} hours.
            
            **Time until reset:** {hours} hours and {minutes} minutes
            
            Please wait before making more requests or contact support for additional quota.
            """)
            return True
    return False


def search_youtube(query):
    try:
        import requests
        import re
        from urllib.parse import quote
        
        # Format the search query
        search_query = quote(query)
        search_url = f"https://www.youtube.com/results?search_query={search_query}"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(search_url, headers=headers, timeout=10)
        
        # Extract video ID from the response
        video_pattern = r'"videoId":"([^"]+)"'
        matches = re.findall(video_pattern, response.text)
        
        if matches:
            video_id = matches[0]
            video_link = f"https://www.youtube.com/watch?v={video_id}"
            return video_link
        else:
            return "Video not found."
            
    except Exception as e:
        return f"Error searching YouTube: {str(e)}"


def search_pdf(query):
    try:
        pdf_query = f"filetype:pdf {query}"
        for j in search(pdf_query, tld="co.in", num=1, stop=1, pause=2):
            if ".pdf" in j:
                return f"PDF found: {j}"
        return "PDF not found."
    except Exception as e:
        return f"Error searching for PDF: {str(e)}"

def search_ppt(query):
    try:
        ppt_query = f"filetype:ppt OR filetype:pptx {query}"
        for j in search(ppt_query, tld="co.in", num=1, stop=1, pause=2):
            if ".ppt" in j or ".pptx" in j:
                return f"PPT found: {j}"
        return "PPT not found."
    except Exception as e:
        return f"Error searching for PPT: {str(e)}"

def search_blogs(query):
    """
    Search for security blogs and articles related to the given concept
    """
    try:
        # Target popular security blog sites and general blog content
        blog_sites = [
            "site:medium.com",
            "site:dev.to", 
            "site:hackernoon.com",
            "site:infosec-handbook.eu",
            "site:pentestlab.blog",
            "site:portswigger.net/research",
            "site:blog.securelayer7.net",
            "site:blog.checkpoint.com",
            "site:research.checkpoint.com"
        ]
        
        # Try searching with different blog site constraints
        for site in blog_sites:
            blog_query = f"{query} {site}"
            try:
                for j in search(blog_query, tld="co.in", num=1, stop=1, pause=2):
                    return f"Blog found: {j}"
            except:
                continue
        
        # If no specific site results, try general blog search
        general_blog_query = f"{query} blog tutorial cybersecurity"
        for j in search(general_blog_query, tld="co.in", num=1, stop=1, pause=2):
            # Filter for likely blog URLs
            if any(keyword in j.lower() for keyword in ['blog', 'medium.com', 'dev.to', 'hackernoon', 'tutorial', 'guide']):
                return f"Blog found: {j}"
        
        return "Blog not found."
        
    except Exception as e:
        return f"Error searching for blogs: {str(e)}"


def search_poc_reports(query):
    """
    Search for Proof of Concept (PoC) exploits and vulnerability disclosure reports
    """
    try:
        # Target sites known for  research, PoCs, and vulnerability disclosures
        poc_sites = [
            "site:github.com",
            "site:exploit-db.com",
            "site:cve.mitre.org",
            "site:nvd.nist.gov",
            "site:packetstorm.com",
            "site:seclists.org",
            "site:focus.com",
            "site:rapid7.com/db",
            "site:vulners.com",
            "site:zerodayinitiative.com"
        ]
        
        # PoC-specific search terms
        poc_terms = ["exploit", "proof of concept", "PoC", "vulnerability", "CVE", "disclosure"]
        
        # Try searching with PoC sites and terms
        for site in poc_sites:
            for term in poc_terms:
                poc_query = f"{query} {term} {site}"
                try:
                    for j in search(poc_query, tld="co.in", num=1, stop=1, pause=2):
                        return f"PoC/Report found: {j}"
                except:
                    continue
        
        # Try general PoC search without site restrictions
        general_poc_queries = [
            f"{query} exploit proof of concept",
            f"{query} vulnerability disclosure",
            f"{query} CVE exploit",
            f"{query}  advisory",
            f"{query} PoC github"
        ]
        
        for poc_query in general_poc_queries:
            try:
                for j in search(poc_query, tld="co.in", num=1, stop=1, pause=2):
                    # Filter for likely PoC/vulnerability report URLs
                    if any(keyword in j.lower() for keyword in ['exploit', 'cve', 'vulnerability', 'poc', 'github.com', '', 'advisory']):
                        return f"PoC/Report found: {j}"
            except:
                continue
        
        return "PoC/Report not found."
        
    except Exception as e:
        return f"Error searching for PoC/Reports: {str(e)}"


# Usage example functions (optional helper functions)
def search_vulnerability_databases(query):
    """
    Specifically search vulnerability databases for CVE information
    """
    try:
        vuln_db_sites = [
            "site:cve.mitre.org",
            "site:nvd.nist.gov", 
            "site:cvedetails.com",
            "site:vuldb.com"
        ]
        
        for site in vuln_db_sites:
            vuln_query = f"{query} CVE {site}"
            try:
                for j in search(vuln_query, tld="co.in", num=1, stop=1, pause=2):
                    return f"Vulnerability DB found: {j}"
            except:
                continue
                
        return "Vulnerability database entry not found."
        
    except Exception as e:
        return f"Error searching vulnerability databases: {str(e)}"


def search__advisories(query):
    """
    Search for official  advisories and vendor disclosures
    """
    try:
        advisory_sites = [
            "site:security.microsoft.com",
            "site:support.apple.com/security",
            "site:chromereleases.googleblog.com",
            "site:mozilla.org/security",
            "site:us-cert.cisa.gov",
            "site:cert.org"
        ]
        
        for site in advisory_sites:
            advisory_query = f"{query} security advisory {site}"
            try:
                for j in search(advisory_query, tld="co.in", num=1, stop=1, pause=2):
                    return f"Security Advisory found: {j}"
            except:
                continue
                
        return "Security advisory not found."
        
    except Exception as e:
        return f"Error searching security advisories: {str(e)}"
class SecurityAnalysisApp:
    def __init__(self):
        self.chat_model = ChatGroq(
            groq_api_key=st.secrets["groq_api_key"],
            model_name="qwen-qwq-32b",
            temperature=0.7,
            max_tokens=None
        )
        
        self.analysis_template = PromptTemplate(
            input_variables=["code_chunk"],
            template="""Analyze the following code for malicious indicators and security concerns.

Code to analyze:
{code_chunk}

Provide a comprehensive analysis with a brief summary of findings and details for each category. Return your analysis in this exact JSON structure:
{{
    "summary": ["Brief overview of key findings and potential security implications"],
    "sections": {{
        "code_obfuscation_techniques": {{
            "findings": [],
            "description": "Brief explanation of identified obfuscation techniques and their implications"
        }},
        "suspicious_api_calls": {{
            "findings": [],
            "description": "Overview of concerning API calls and their potential security impact"
        }},
        "anti_analysis_mechanisms": {{
            "findings": [],
            "description": "Summary of detected anti-analysis features and their significance"
        }},
        "network_communication_patterns": {{
            "findings": [],
            "description": "Analysis of network-related code patterns and security concerns"
        }},
        "file_system_operations": {{
            "findings": [],
            "description": "Evaluation of file system interactions and associated risks"
        }},
        "potential_payload_analysis": {{
            "findings": [],
            "description": "Assessment of potential malicious payloads and their characteristics"
        }}
    }}
}}

Requirements:
1. Ensure each field is populated with meaningful content
2. Include a clear summary of overall findings
3. Provide a brief description for each section
4. List specific findings as bullet points in the findings arrays
5. Use "None identified" in findings array if no indicators are found
6. Keep descriptions concise and focused on security implications

Your response should be ONLY the JSON object with no additional text."""
        )
        
        self.binary_analysis_template = PromptTemplate(
            input_variables=["strings_chunk"],
            template="""Analyze the following strings extracted from a binary file for malicious indicators and security concerns.

Extracted strings:
{strings_chunk}

Provide a comprehensive analysis with a brief summary of findings and details for each category. Return your analysis in this exact JSON structure:
{{
    "summary": ["Brief overview of key findings and potential security implications"],
    "sections": {{
        "suspicious_strings": {{
            "findings": [],
            "description": "Brief explanation of identified suspicious strings and their implications"
        }},
        "command_and_control_indicators": {{
            "findings": [],
            "description": "Overview of potential C2 indicators like URLs, IPs, or domain patterns"
        }},
        "anti_analysis_indicators": {{
            "findings": [],
            "description": "Summary of strings suggesting anti-analysis capabilities"
        }},
        "network_related_strings": {{
            "findings": [],
            "description": "Analysis of network-related strings and security concerns"
        }},
        "file_system_indicators": {{
            "findings": [],
            "description": "Evaluation of file system related strings and associated risks"
        }},
        "potential_malware_functionality": {{
            "findings": [],
            "description": "Assessment of strings indicating malicious functionality"
        }}
    }}
}}

Requirements:
1. Ensure each field is populated with meaningful content
2. Include a clear summary of overall findings
3. Provide a brief description for each section
4. List specific findings as bullet points in the findings arrays
5. Use "None identified" in findings array if no indicators are found
6. Keep descriptions concise and focused on security implications

Your  should be ONLY the JSON object with no additional text and TRY TO KEEP THE TOKEN SIZE MINIMUM."""
        )
        
        self.analysis_chain = LLMChain(
            llm=self.chat_model,
            prompt=self.analysis_template,
            verbose=True
        )
        
        self.binary_analysis_chain = LLMChain(
            llm=self.chat_model,
            prompt=self.binary_analysis_template,
            verbose=True
        )
        
        self.chat_memory = ConversationBufferMemory()
        self.conversation = ConversationChain(
            llm=self.chat_model,
            memory=self.chat_memory,
            verbose=True
        )

    def clean_json_response(self, response: str) -> str:
        try:
            start = response.find('{')
            end = response.rfind('}') + 1
            if start != -1 and end != 0:
                response = response[start:end]
            response = response.replace('```json', '').replace('```', '')
            return response.strip()
        except Exception:
            return response

    def analyze_chunk(self, chunk: str, is_binary: bool = False) -> Dict:
        try:
            if is_binary:
                response = self.binary_analysis_chain.predict(strings_chunk=chunk)
            else:
                response = self.analysis_chain.predict(code_chunk=chunk)
                
            cleaned_response = self.clean_json_response(response)
            
            try:
                return json.loads(cleaned_response)
            except json.JSONDecodeError as je:
                return self._create_error_analysis("JSON parsing failed", str(je), is_binary)
                
        except Exception as e:
            return self._create_error_analysis("Analysis failed", str(e), is_binary)

    def _create_error_analysis(self, error_type: str, details: str, is_binary: bool = False) -> Dict:
        if is_binary:
            sections = [
                "suspicious_strings",
                "command_and_control_indicators",
                "anti_analysis_indicators",
                "network_related_strings",
                "file_system_indicators",
                "potential_malware_functionality"
            ]
        else:
            sections = [
                "code_obfuscation_techniques",
                "suspicious_api_calls",
                "anti_analysis_mechanisms",
                "network_communication_patterns",
                "file_system_operations",
                "potential_payload_analysis"
            ]
            
        return {
            "error": f"{error_type}: {details}",
            "summary": ["Analysis failed - " + error_type],
            "sections": {section: {
                "findings": ["Analysis failed"],
                "description": "Analysis failed due to technical error"
            } for section in sections}
        }

    def split_code_in_chunks(self, content: str, chunk_size: int = 12800) -> List[str]:
        return [content[i:i + chunk_size] for i in range(0, len(content), chunk_size)]

    def analyze_code(self, code_content: str) -> Dict:
        chunks = self.split_code_in_chunks(code_content)
        analyses = []
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, chunk in enumerate(chunks, 1):
            status_text.text(f"Analyzing chunk {i}/{len(chunks)}...")
            analysis = self.analyze_chunk(chunk)
            analyses.append(analysis)
            progress_bar.progress(i/len(chunks))
        
        status_text.text("Analysis complete!")
        progress_bar.empty()
        
        return self.combine_analyses(analyses)
        

    def extract_strings_from_binary(self, binary_data: bytes, min_length: int = 10) -> str:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.exe') as temp_file:
            temp_file.write(binary_data)
            temp_file_path = temp_file.name

        try:
            try:
                # Use 'strings' with a length filter if available
                result = subprocess.run(['strings', '-n', str(min_length), temp_file_path],
                                        capture_output=True, text=True, check=True)
                extracted_strings = result.stdout
            except (subprocess.SubprocessError, FileNotFoundError):
                # Fallback: Extract printable ASCII strings manually
                extracted_strings = self._extract_strings_manually(binary_data, min_length)

            # Filter output (optional, e.g., extract only strings containing certain keywords)
            filtered_strings = "\n".join(
                line for line in extracted_strings.splitlines()
                if re.search(r'[a-zA-Z0-9_]', line)  # Ensure meaningful content
            )
            
            return " ".join(filtered_strings.splitlines())
        finally:
            try:
                os.unlink(temp_file_path)
            except:
                pass

    
    def _extract_strings_manually(self, binary_data: bytes, min_length: int = 4) -> str:
        strings = []
        current_string = ""
        
        for byte in binary_data:
            # Check if byte is printable ASCII
            if 32 <= byte <= 126:  # Printable ASCII range
                current_string += chr(byte)
            else:
                if len(current_string) >= min_length:
                    strings.append(current_string)
                current_string = ""
                
        # Add the last string if it meets the minimum length
        if len(current_string) >= min_length:
            strings.append(current_string)
            
        return "\n".join(strings)
    
    def analyze_binary(self, binary_data: bytes) -> Dict:
        strings_content = self.extract_strings_from_binary(binary_data)
        chunks = self.split_code_in_chunks(strings_content)
        analyses = []
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, chunk in enumerate(chunks, 1):
            status_text.text(f"Analyzing binary chunk {i}/{len(chunks)}...")
            analysis = self.analyze_chunk(chunk, is_binary=True)
            analyses.append(analysis)
            progress_bar.progress(i/len(chunks))
        
        status_text.text("Binary analysis complete!")
        progress_bar.empty()
        
        return self.combine_analyses(analyses)

    def combine_analyses(self, analyses: List[Dict]) -> Dict:
        # Get all section keys from the first analysis
        if not analyses:
            return {
                "summary": ["No analysis results available"],
                "sections": {},
                "errors": ["No analyses performed"]
            }
            
        # Initialize combined structure with all possible sections
        combined = {
            "summary": set(),
            "sections": {},
            "errors": []
        }
        
        # Initialize sections based on first analysis
        if "sections" in analyses[0]:
            for section in analyses[0]["sections"]:
                combined["sections"][section] = {"findings": set(), "description": ""}
        
        for analysis in analyses:
            if "error" in analysis:
                combined["errors"].append(analysis["error"])
            
            if "summary" in analysis:
                combined["summary"].update(analysis["summary"])
            
            if "sections" in analysis:
                for section, content in analysis["sections"].items():
                    if section not in combined["sections"]:
                        combined["sections"][section] = {"findings": set(), "description": ""}
                    
                    if "findings" in content:
                        combined["sections"][section]["findings"].update(content["findings"])
                    
                    if content.get("description") and not combined["sections"][section]["description"]:
                        combined["sections"][section]["description"] = content["description"]

        result = {
            "summary": list(combined["summary"]),
            "sections": {},
            "errors": combined["errors"]
        }

        for section, content in combined["sections"].items():
            findings = list(content["findings"])
            if len(findings) > 1 and "Analysis failed" in findings:
                findings.remove("Analysis failed")
            
            result["sections"][section] = {
                "findings": findings,
                "description": content["description"] or "No significant findings in this category"
            }

        return result

    def create_analysis_report(self, analysis_results: Dict, title: str = "Security Analysis Report") -> str:
        document = Document()
        document.add_heading(title, 0)
        
        document.add_heading("Executive Summary", level=1)
        for summary_point in analysis_results.get("summary", ["No summary available"]):
            document.add_paragraph(summary_point, style='Body Text')
        
        for section_name, content in analysis_results.get("sections", {}).items():
            heading_text = section_name.replace('_', ' ').title()
            document.add_heading(heading_text, level=1)
            
            if content.get("description"):
                document.add_paragraph(content["description"], style='Body Text')
            
            if content.get("findings"):
                document.add_heading("Findings:", level=2)
                for finding in content["findings"]:
                    if finding != "None identified":
                        document.add_paragraph(f"• {finding}", style='List Bullet')
                    else:
                        document.add_paragraph("No specific issues identified in this category.", style='Body Text')
        
        if analysis_results.get("errors"):
            document.add_heading("Analysis Errors", level=1)
            for error in analysis_results["errors"]:
                document.add_paragraph(f"• {error}", style='List Bullet')
        
        report_filename = f"security_analysis_report_{os.getpid()}.docx"
        document.save(report_filename)
        return report_filename

    def get_chat_response(self, user_input: str) -> str:
        return self.conversation.predict(input=user_input+"Response should be short and crisp")

def display_analysis_results(analysis: Dict):
    st.header("Executive Summary")
    for summary_point in analysis.get("summary", ["No summary available"]):
        st.write(summary_point)
    st.divider()

    if analysis.get("errors"):
        st.error("Analysis Errors")
        for error in analysis["errors"]:
            st.write(f"• {error}")
        st.divider()

    for section_name, content in analysis.get("sections", {}).items():
        st.subheader(section_name.replace('_', ' ').title())
        
        if content.get("description"):
            st.write(content["description"])
        
        if content.get("findings"):
            st.write("Findings:")
            for finding in content["findings"]:
                st.write(f"• {finding}")
        st.divider()



class CyberSecurityAssistant:
    def __init__(self, groq_api_key):
        self.llm = ChatGroq(
            model="llama-3.3-70b-versatile",
            temperature=0.3,
            max_tokens=None,
            timeout=None,
            max_retries=2,
            groq_api_key=groq_api_key,
        )
        
        self.prompt = ChatPromptTemplate.from_messages([
            (
                "system",
                """You are GamkersGPT, an ethical hacking trainer who helps students understand cybersecurity concepts and practices. 
                You assist students by generating explanations, code snippets, suggesting tools, or providing step-by-step instructions based on the task and scenario provided.
                
                All information you provide should be for educational purposes only. Emphasize ethical practices and legal boundaries in all responses.
                Always remind users that any security testing should only be performed on systems they own or have explicit permission to test.
                
                Your expertise includes:
                - Network security and penetration testing
                - Web application security
                - Security tools and frameworks
                - Vulnerability assessment techniques
                - Defensive security practices and countermeasures
                - Common security tools commands and usage
                - Google dorking techniques
                - CVE understanding and analysis
                - Latest cybersecurity trends and news
                
                When providing technical guidance:
                1. Explain the concepts first
                2. Outline the methodology
                3. Share relevant code or commands with explanations
                4. Highlight security considerations and ethical implications
                """
            ),
            ("human", "{query}")
        ])
        
        # Create the processing chain
        self.chain = self.prompt | self.llm
        
        # Store conversation history
        self.conversation_history = []
        
    def query(self, user_query):
        """Process a user query and return a response"""
        try:
            # Add query to conversation history
            self.conversation_history.append({"role": "user", "content": user_query})
            
            # Process the query
            response = self.chain.invoke({"query": user_query})
            response_text = response.content
            
            # Add response to conversation history
            self.conversation_history.append({"role": "assistant", "content": response_text})
            
            return response_text
        except Exception as e:
            return f"Error processing your query: {str(e)}"
            
    def search_cve(self, cve_query, limit=5):
        """Search for CVE information"""
        try:
            # Format to search for specific CVE ID or general search
            if re.match(r'CVE-\d{4}-\d{4,}', cve_query, re.IGNORECASE):
                cve_id = cve_query.upper()
                url = f"https://services.nvd.nist.gov/rest/json/cves/2.0?cveId={cve_id}"
            else:
                # General keyword search
                url = f"https://services.nvd.nist.gov/rest/json/cves/2.0?keywordSearch={cve_query}&resultsPerPage={limit}"
            
            headers = {
                "User-Agent": "GamkersGPT Educational Tool/1.0"
            }
            
            response = requests.get(url, headers=headers)
            
            if response.status_code == 200:
                data = response.json()
                
                # Process the response to extract relevant information
                results = []
                if 'vulnerabilities' in data:
                    for vuln in data['vulnerabilities'][:limit]:
                        cve_item = vuln['cve']
                        cve_id = cve_item['id']
                        
                        # Get description
                        descriptions = cve_item.get('descriptions', [])
                        description = next((desc['value'] for desc in descriptions if desc['lang'] == 'en'), "No description available")
                        
                        # Get severity if available
                        metrics = cve_item.get('metrics', {})
                        cvss_data = metrics.get('cvssMetricV31', [{}])[0] if 'cvssMetricV31' in metrics else metrics.get('cvssMetricV30', [{}])[0] if 'cvssMetricV30' in metrics else {}
                        base_score = cvss_data.get('cvssData', {}).get('baseScore', "N/A") if cvss_data else "N/A"
                        severity = cvss_data.get('cvssData', {}).get('baseSeverity', "N/A") if cvss_data else "N/A"
                        
                        # Get published and last modified dates
                        published = cve_item.get('published', "N/A")
                        last_modified = cve_item.get('lastModified', "N/A")
                        
                        # Format dates if they exist
                        if published != "N/A":
                            published = published.split('T')[0]  # Just take the date part
                        if last_modified != "N/A":
                            last_modified = last_modified.split('T')[0]  # Just take the date part
                        
                        results.append({
                            'cve_id': cve_id,
                            'description': description,
                            'severity': severity,
                            'base_score': base_score,
                            'published': published,
                            'last_modified': last_modified,
                            'references': cve_item.get('references', [])
                        })
                
                if not results:
                    return f"No CVE entries found for '{cve_query}'"
                
                return results
            else:
                return f"Error searching CVE database: {response.status_code} - {response.text}"
                
        except Exception as e:
            return f"Error searching CVE database: {str(e)}"
    
    def get_hacker_news(self, limit=5):
        """Fetch and summarize latest cybersecurity news from The Hacker News using requests and xml.etree"""
        try:
            feed_url = "https://feeds.feedburner.com/TheHackersNews"
            response = requests.get(feed_url)
            
            if response.status_code != 200:
                return "Unable to fetch cybersecurity news at this time."
            
            root = ET.fromstring(response.content)
            ns = {'atom': 'http://www.w3.org/2005/Atom', 'rss': 'http://www.w3.org/2005/Atom'}  # default Atom ns

            news_items = []
            entries = root.findall(".//item")
            
            for entry in entries[:limit]:
                title = entry.findtext("title")
                link = entry.findtext("link")
                published = entry.findtext("pubDate")
                summary = entry.findtext("description")

                # Clean summary from HTML tags
                summary = re.sub(r'<.*?>', '', summary)

                # Date formatting
                try:
                    pub_date = datetime.strptime(published, "%a, %d %b %Y %H:%M:%S %z")
                    published_formatted = pub_date.strftime("%Y-%m-%d")
                except:
                    published_formatted = published

                # Generate LLM summary
                summary_prompt = f"Summarize this cybersecurity news article in 2-3 sentences, highlighting key security implications:\n\nTitle: {title}\n\nContent: {summary}"

                # Process with LLM
                summary_response = self.llm.invoke(summary_prompt)
                llm_summary = summary_response.content

                news_items.append({
                    'title': title,
                    'published': published_formatted,
                    'link': link,
                    'summary': llm_summary
                })

        
            return news_items

        except Exception as e:
            return f"Error fetching cybersecurity news: {str(e)}"
    
    def generate_tool_commands(self, tool_name, task_description):
        """Generate commands for common security tools based on the task"""
        try:
            prompt = f"""Generate practical command examples for using the security tool '{tool_name}' to accomplish the following task: {task_description}.
            
            Format your response as follows:
            1. Brief introduction to what {tool_name} is and how it's useful for this task
            2. 3-5 practical command examples with syntax highlighting
            3. Explanation of key parameters and options
            4. Common usage scenarios
            5. Security considerations and ethical usage reminder
            
            Be detailed and provide real commands that would actually work in a security testing environment.
            """
            
            response = self.llm.invoke(prompt)
            return response.content
            
        except Exception as e:
            return f"Error generating tool commands: {str(e)}"
    
    def generate_google_dorks(self, target_type, objective):
        """Generate Google dorking queries for specific objectives"""
        try:
            prompt = f"""Generate 5-7 effective Google dorking queries that could be used to find {target_type} with the objective of {objective}.
            
            For each query:
            1. Show the exact Google dork syntax
            2. Explain what the query does and why it's effective
            3. Note any specific operators used
            
            Include a brief introduction explaining Google dorking, its legitimate security applications, and ethical considerations.
            
            Reminder: These are for educational purposes and security assessment only.
            """
            
            response = self.llm.invoke(prompt)
            return response.content
            
        except Exception as e:
            return f"Error generating Google dorks: {str(e)}"


    def analyze_network_traffic(self, pcap_data=None, sample_type="http"):
        """Analyze network traffic patterns or provided PCAP samples
        
        Args:
            pcap_data: Optional base64 encoded PCAP data or description of traffic
            sample_type: Type of traffic to analyze (http, dns, tcp, etc.)
        """
        try:
            prompt = f"""Analyze the following {sample_type} network traffic scenario and provide security insights:
            
            {pcap_data if pcap_data else f'Provide a detailed analysis of common {sample_type} traffic patterns and security implications'}
            
            In your analysis include:
            1. Key indicators to look for in this type of traffic
            2. Common security issues or attack patterns in {sample_type} traffic
            3. How to identify anomalies or malicious activities
            4. Tools and techniques for monitoring this traffic type
            5. Recommended security controls and best practices
            """
            
            response = self.llm.invoke(prompt)
            return response.content
        except Exception as e:
            return f"Error analyzing network traffic: {str(e)}"
    
    def generate_security_policy(self, organization_type, focus_area):
        """Generate security policy templates based on organization type and focus area
        
        Args:
            organization_type: Type of organization (healthcare, finance, education, etc.)
            focus_area: Specific security focus (network, data, access control, etc.)
        """
        try:
            prompt = f"""Generate a comprehensive security policy template for a {organization_type} organization, 
            focusing specifically on {focus_area} security.
            
            The policy should include:
            1. Purpose and scope
            2. Roles and responsibilities
            3. Specific policy statements and controls
            4. Compliance requirements relevant to {organization_type} organizations
            5. Implementation guidance
            6. Monitoring and enforcement mechanisms
            7. Review and update procedures
            
            Format the policy in a professional, structured manner suitable for organizational use.
            """
            
            response = self.llm.invoke(prompt)
            return response.content
        except Exception as e:
            return f"Error generating security policy: {str(e)}"
    
    def analyze_attack_scenario(self, attack_vector, target_system):
        """Analyze a security attack scenario and provide defense strategies
        
        Args:
            attack_vector: Type of attack (phishing, ransomware, SQLi, etc.)
            target_system: Type of system being targeted (web server, API, etc.)
        """
        try:
            prompt = f"""Analyze the following attack scenario and provide a comprehensive defense strategy:
            
            Attack Vector: {attack_vector}
            Target System: {target_system}
            
            Your analysis should include:
            1. Attack methodology and typical progression
            2. Initial compromise techniques
            3. Potential lateral movement and persistence mechanisms
            4. Key vulnerabilities exploited
            5. Detection strategies and indicators of compromise
            6. Prevention measures (before attack)
            7. Mitigation strategies (during attack)
            8. Recovery procedures (after attack)
            9. Relevant security tools and frameworks for defense
            
            Keep your recommendations ethical and focused on defensive security.
            """
            
            response = self.llm.invoke(prompt)
            return response.content
        except Exception as e:
            return f"Error analyzing attack scenario: {str(e)}"
    
    def decode_or_encode_data(self, input_data, operation, format_type):
        """Encode or decode data in various formats
        
        Args:
            input_data: The data to encode/decode
            operation: "encode" or "decode"
            format_type: base64, hex, url, etc.
        """
        try:
            import base64
            import urllib.parse
            import codecs
            
            # Remove whitespace for consistent processing
            input_data = input_data.strip()
            
            result = ""
            explanation = ""
            
            if operation == "encode":
                if format_type == "base64":
                    result = base64.b64encode(input_data.encode()).decode()
                    explanation = "Converted text to Base64 encoded string"
                elif format_type == "hex":
                    result = input_data.encode().hex()
                    explanation = "Converted text to hexadecimal representation"
                elif format_type == "url":
                    result = urllib.parse.quote(input_data)
                    explanation = "URL encoded the input string"
                elif format_type == "binary":
                    result = ' '.join(format(ord(c), '08b') for c in input_data)
                    explanation = "Converted text to binary representation"
                else:
                    return f"Unsupported encoding format: {format_type}"
            
            elif operation == "decode":
                try:
                    if format_type == "base64":
                        result = base64.b64decode(input_data).decode()
                        explanation = "Decoded Base64 string to text"
                    elif format_type == "hex":
                        # Handle both with and without spaces
                        clean_hex = input_data.replace(" ", "")
                        result = bytes.fromhex(clean_hex).decode()
                        explanation = "Converted hexadecimal to text"
                    elif format_type == "url":
                        result = urllib.parse.unquote(input_data)
                        explanation = "URL decoded the input string"
                    elif format_type == "binary":
                        # Handle binary with spaces
                        binary_values = input_data.split()
                        result = ''.join(chr(int(bin_val, 2)) for bin_val in binary_values)
                        explanation = "Converted binary to text"
                    else:
                        return f"Unsupported decoding format: {format_type}"
                except Exception as decode_error:
                    return f"Error decoding {format_type}: {str(decode_error)}. Please check that your input is valid {format_type} format."
            else:
                return f"Unsupported operation: {operation}. Use 'encode' or 'decode'."
                
            return {
                "result": result,
                "explanation": explanation,
                "operation": operation,
                "format": format_type
            }
            
        except Exception as e:
            return f"Error processing data: {str(e)}"
    
    def hash_analyzer(self, hash_value):
        """Analyze a hash value to determine type and provide information
        
        Args:
            hash_value: The hash string to analyze
        """
        try:
            hash_value = hash_value.strip()
            hash_length = len(hash_value)
            
            # Common hash types and their characteristics
            hash_types = {
                32: {
                    "name": "MD5",
                    "strength": "Weak - vulnerable to collision attacks",
                    "pattern": r"^[a-fA-F0-9]{32}$"
                },
                40: {
                    "name": "SHA-1",
                    "strength": "Weak - vulnerable to collision attacks",
                    "pattern": r"^[a-fA-F0-9]{40}$"
                },
                64: {
                    "name": "SHA-256",
                    "strength": "Strong - currently considered secure",
                    "pattern": r"^[a-fA-F0-9]{64}$"
                },
                96: {
                    "name": "SHA-384",
                    "strength": "Strong - currently considered secure",
                    "pattern": r"^[a-fA-F0-9]{96}$"
                },
                128: {
                    "name": "SHA-512",
                    "strength": "Very Strong - currently considered secure",
                    "pattern": r"^[a-fA-F0-9]{128}$"
                }
            }
            
            import re
            
            # Check if the hash matches the pattern for its length
            if hash_length in hash_types:
                hash_info = hash_types[hash_length]
                if re.match(hash_info["pattern"], hash_value):
                    result = {
                        "hash_type": hash_info["name"],
                        "length": hash_length,
                        "strength": hash_info["strength"],
                        "valid_format": True,
                        "additional_info": f"This appears to be a valid {hash_info['name']} hash."
                    }
                else:
                    result = {
                        "hash_type": "Unknown",
                        "length": hash_length,
                        "strength": "Unknown",
                        "valid_format": False,
                        "additional_info": "This string has the right length for some hash types but contains invalid characters."
                    }
            else:
                # Check for other hash types or formats
                if re.match(r"^\$2[ayb]\$[0-9]{2}\$[A-Za-z0-9./]{53}$", hash_value):
                    result = {
                        "hash_type": "bcrypt",
                        "length": hash_length,
                        "strength": "Strong - designed to be slow for password hashing",
                        "valid_format": True,
                        "additional_info": "This appears to be a bcrypt password hash."
                    }
                elif re.match(r"^\$6\$[a-zA-Z0-9./]{8,16}\$[a-zA-Z0-9./]{86}$", hash_value):
                    result = {
                        "hash_type": "SHA-512 crypt",
                        "length": hash_length,
                        "strength": "Strong",
                        "valid_format": True,
                        "additional_info": "This appears to be a SHA-512 crypt hash used in Linux/Unix systems."
                    }
                else:
                    result = {
                        "hash_type": "Unknown",
                        "length": hash_length,
                        "strength": "Unknown",
                        "valid_format": "Unknown",
                        "additional_info": "This string doesn't match common hash formats. It may be a custom hash format, encrypted data, or not a hash at all."
                    }
            
            return result
            
        except Exception as e:
            return f"Error analyzing hash: {str(e)}"
    
    def vulnerability_assessment(self, system_description):
        """Generate a vulnerability assessment based on a system description
        
        Args:
            system_description: Description of the system to assess
        """
        try:
            prompt = f"""Perform a vulnerability assessment for the following system:

            {system_description}
            
            In your assessment, include:
            1. Potential vulnerability categories based on the system description
            2. Specific vulnerabilities likely to affect this system
            3. OWASP Top 10 or SANS Top 25 vulnerabilities that apply
            4. Assessment methodology recommendations
            5. Testing approaches for each vulnerability category
            6. Mitigation strategies prioritized by risk
            7. Recommended security tools for testing
            
            Present this as a professional vulnerability assessment report.
            """
            
            response = self.llm.invoke(prompt)
            return response.content
        except Exception as e:
            return f"Error generating vulnerability assessment: {str(e)}"
    
    def generate_incident_response_plan(self, incident_type, organization_size):
        """Generate an incident response plan template
        
        Args:
            incident_type: Type of security incident (data breach, ransomware, etc.)
            organization_size: Size of the organization (small, medium, large)
        """
        try:
            prompt = f"""Generate a comprehensive incident response plan for a {organization_size} organization dealing with a {incident_type} incident.
            
            The plan should include:
            1. Incident response team structure and roles
            2. Preparation measures specific to {incident_type}
            3. Detection and identification procedures
            4. Containment strategies
            5. Eradication and recovery steps
            6. Post-incident analysis methodology
            7. Communication plan (internal and external)
            8. Legal and compliance considerations
            9. Documentation requirements
            10. Training recommendations
            
            Format this as a complete incident response playbook that could be implemented by a {organization_size} organization.
            """
            
            response = self.llm.invoke(prompt)
            return response.content
        except Exception as e:
            return f"Error generating incident response plan: {str(e)}"
    
    def analyze_code_security(self, code_snippet, language):
        """Analyze code for security vulnerabilities
        
        Args:
            code_snippet: The code to analyze
            language: Programming language of the code
        """
        try:
            prompt = f"""Perform a security code review on the following {language} code:
            
            ```{language}
            {code_snippet}
            ```
            
            In your analysis:
            1. Identify security vulnerabilities or weaknesses
            2. Categorize each issue by type (e.g., injection, authentication, etc.)
            3. Assess the risk level of each issue (Critical, High, Medium, Low)
            4. Provide specific line references where problems occur
            5. Explain why each issue is a security concern
            6. Recommend secure coding fixes with example code
            7. Suggest secure coding practices relevant to {language}
            
            Format your response as a professional security code review report.
            """
            
            response = self.llm.invoke(prompt)
            return response.content
        except Exception as e:
            return f"Error analyzing code security: {str(e)}"
    
    def generate_ctf_challenge(self, difficulty, category):
        """Generate a Capture The Flag (CTF) challenge for practice
        
        Args:
            difficulty: Difficulty level (easy, medium, hard)
            category: Challenge category (web, crypto, forensics, etc.)
        """
        try:
            prompt = f"""Create a {difficulty} level {category} Capture The Flag (CTF) challenge for cybersecurity training.
            
            The challenge should include:
            1. A compelling scenario or background story
            2. Detailed challenge description
            3. Any necessary setup instructions or environment details
            4. Files or code that would need to be created (described in detail)
            5. Step-by-step solution guide (for the instructor)
            6. Hints that could be provided to participants (at least 3)
            7. Learning objectives and security concepts demonstrated
            8. Suggested point value for a CTF competition
            
            The challenge should be realistic, educational, and engaging while adhering to ethical standards.
            """
            
            response = self.llm.invoke(prompt)
            return response.content
        except Exception as e:
            return f"Error generating CTF challenge: {str(e)}"
    
    def explain_security_concept(self, concept_name):
        """Provide detailed explanation of security concepts with practical examples
        
        Args:
            concept_name: The security concept to explain
        """
        try:
            prompt = f"""Provide a comprehensive explanation of the security concept: {concept_name}
            
            In your explanation, include:
            1. Definition and core principles
            2. Historical context and development
            3. How this concept works technically
            4. Real-world applications and implementations
            5. Common attacks or vulnerabilities related to this concept
            6. Defensive strategies and best practices
            7. Practical examples that illustrate the concept
            8. Tools or frameworks associated with this concept
            9. Future trends or developments
            10. Resources for further learning
            
            Make the explanation accessible but technically accurate, suitable for cybersecurity training *NOTE EXPLAIN EVERYTHING IN SHORT*.
            """
            
            response = self.llm.invoke(prompt)
            return response.content
        except Exception as e:
            return f"Error explaining security concept: {str(e)}"
            
    def decrypt_or_encrypt_data(self, input_data, operation, algorithm, key=None, iv=None):
        """Encrypt or decrypt data using common cryptographic algorithms
        
        Args:
            input_data: Data to encrypt/decrypt
            operation: "encrypt" or "decrypt"
            algorithm: Algorithm to use (aes, des, etc.)
            key: Encryption/decryption key (will generate if none provided for encrypt)
            iv: Initialization vector (will generate if none provided for encrypt)
        """
        try:
            from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
            from cryptography.hazmat.backends import default_backend
            import os
            import base64
            
            supported_algorithms = {
                "aes": {
                    "key_sizes": [16, 24, 32],  # AES-128, AES-192, AES-256
                    "block_size": 16,
                    "algorithm": algorithms.AES
                },
                "des": {
                    "key_sizes": [8],
                    "block_size": 8,
                    "algorithm": algorithms.TripleDES
                },
                "chacha20": {
                    "key_sizes": [32],
                    "nonce_size": 16,
                    "algorithm": algorithms.ChaCha20
                }
            }
            
            if algorithm.lower() not in supported_algorithms:
                return f"Unsupported algorithm: {algorithm}. Supported algorithms are: {', '.join(supported_algorithms.keys())}"
            
            algo_info = supported_algorithms[algorithm.lower()]
            
            # Handling key
            if operation == "encrypt":
                if key is None:
                    # Generate a random key of the first valid size
                    key = os.urandom(algo_info["key_sizes"][0])
                    key_b64 = base64.b64encode(key).decode()
                else:
                    # Use provided key, ensuring correct size
                    try:
                        key_bytes = base64.b64decode(key)
                        if len(key_bytes) not in algo_info["key_sizes"]:
                            return f"Invalid key size for {algorithm}. Valid sizes (bytes): {algo_info['key_sizes']}"
                        key = key_bytes
                        key_b64 = key
                    except:
                        return "Invalid key format. Key must be base64 encoded."
            else:  # Decrypt
                if key is None:
                    return "Decryption requires a key."
                try:
                    key = base64.b64decode(key)
                    if len(key) not in algo_info["key_sizes"]:
                        return f"Invalid key size for {algorithm}. Valid sizes (bytes): {algo_info['key_sizes']}"
                except:
                    return "Invalid key format. Key must be base64 encoded."
            
            # AES and DES require a mode (CBC, GCM, etc.)
            if algorithm.lower() in ["aes", "des"]:
                # Handle IV for block ciphers
                if operation == "encrypt":
                    if iv is None:
                        iv = os.urandom(algo_info["block_size"])
                        iv_b64 = base64.b64encode(iv).decode()
                    else:
                        try:
                            iv_bytes = base64.b64decode(iv)
                            if len(iv_bytes) != algo_info["block_size"]:
                                return f"Invalid IV size for {algorithm}. Required size: {algo_info['block_size']} bytes"
                            iv = iv_bytes
                            iv_b64 = iv
                        except:
                            return "Invalid IV format. IV must be base64 encoded."
                else:  # Decrypt
                    if iv is None:
                        return f"{algorithm} decryption in CBC mode requires an IV."
                    try:
                        iv = base64.b64decode(iv)
                        if len(iv) != algo_info["block_size"]:
                            return f"Invalid IV size for {algorithm}. Required size: {algo_info['block_size']} bytes"
                    except:
                        return "Invalid IV format. IV must be base64 encoded."
                
                # Create cipher
                cipher = Cipher(
                    algo_info["algorithm"](key),
                    modes.CBC(iv),
                    backend=default_backend()
                )
            
            elif algorithm.lower() == "chacha20":
                if operation == "encrypt":
                    if iv is None:
                        iv = os.urandom(algo_info["nonce_size"])
                        iv_b64 = base64.b64encode(iv).decode()
                    else:
                        try:
                            iv_bytes = base64.b64decode(iv)
                            if len(iv_bytes) != algo_info["nonce_size"]:
                                return f"Invalid nonce size for ChaCha20. Required size: {algo_info['nonce_size']} bytes"
                            iv = iv_bytes
                            iv_b64 = iv
                        except:
                            return "Invalid nonce format. Nonce must be base64 encoded."
                else:  # Decrypt
                    if iv is None:
                        return "ChaCha20 decryption requires a nonce."
                    try:
                        iv = base64.b64decode(iv)
                        if len(iv) != algo_info["nonce_size"]:
                            return f"Invalid nonce size for ChaCha20. Required size: {algo_info['nonce_size']} bytes"
                    except:
                        return "Invalid nonce format. Nonce must be base64 encoded."
                
                # Create cipher
                cipher = Cipher(
                    algo_info["algorithm"](key, iv),
                    mode=None,  # ChaCha20 doesn't use a mode
                    backend=default_backend()
                )
            
            # Perform encryption/decryption
            if operation == "encrypt":
                # Add padding for block ciphers
                if algorithm.lower() in ["aes", "des"]:
                    from cryptography.hazmat.primitives import padding
                    padder = padding.PKCS7(algo_info["block_size"] * 8).padder()
                    data = padder.update(input_data.encode()) + padder.finalize()
                else:
                    data = input_data.encode()
                
                encryptor = cipher.encryptor()
                ciphertext = encryptor.update(data) + encryptor.finalize()
                result = base64.b64encode(ciphertext).decode()
                
                return {
                    "result": result,
                    "key": key_b64 if 'key_b64' in locals() else base64.b64encode(key).decode(),
                    "iv": iv_b64 if 'iv_b64' in locals() else base64.b64encode(iv).decode() if iv else None,
                    "algorithm": algorithm,
                    "operation": "encrypt"
                }
                
            else:  # Decrypt
                try:
                    ciphertext = base64.b64decode(input_data)
                    decryptor = cipher.decryptor()
                    decrypted_data = decryptor.update(ciphertext) + decryptor.finalize()
                    
                    # Remove padding for block ciphers
                    if algorithm.lower() in ["aes", "des"]:
                        from cryptography.hazmat.primitives import padding
                        unpadder = padding.PKCS7(algo_info["block_size"] * 8).unpadder()
                        decrypted_data = unpadder.update(decrypted_data) + unpadder.finalize()
                    
                    return {
                        "result": decrypted_data.decode(),
                        "algorithm": algorithm,
                        "operation": "decrypt"
                    }
                except Exception as decrypt_error:
                    return f"Decryption error: {str(decrypt_error)}. Check that your key, IV, and ciphertext are correct."
            
        except Exception as e:
            return f"Error in cryptographic operation: {str(e)}"
    
    
def main():

    initialize_rate_limiter()
    
    # Display rate limit info in sidebar
    display_rate_limit_info()

    if "app" not in st.session_state:
        st.session_state.app = SecurityAnalysisApp()
        st.session_state.messages = []
    
    st.set_page_config(
        page_title="GamkersGPT AI", 
        page_icon="🔐", 
        layout="wide"
    )

    st.markdown("""
        <style>
        /* Base styles and typography */
        body {
            background-color: linear-gradient(to bottom, rgb(3, 7, 18), rgb(17, 24, 39));
            color: #e2e8f0;
            font-family: 'Inter', sans-serif;
        }

        /* Main header styling */
        .main-header {
            padding: 1.5rem;
            background-color: linear-gradient(to bottom, rgb(3, 7, 18), rgb(17, 24, 39));
            border-radius: 12px;
            margin-bottom: 1.5rem;
            border: 1px solid rgba(148, 163, 184, 0.1);
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
        }

        /* Status indicator */
        .status-indicator {
            width: 12px;
            height: 12px;
            background-color: #10b981;
            border-radius: 50%;
            display: inline-block;
            margin-right: 8px;
            box-shadow: 0 0 0 rgba(16, 185, 129, 0.4);
            animation: pulse 2s infinite;
        }

        @keyframes pulse {
            0% {
                box-shadow: 0 0 0 0 rgba(16, 185, 129, 0.4);
            }
            70% {
                box-shadow: 0 0 0 10px rgba(16, 185, 129, 0);
            }
            100% {
                box-shadow: 0 0 0 0 rgba(16, 185, 129, 0);
            }
        }

        .glow-text {
            color: #10b981;
            font-weight: 500;
            text-shadow: 0 0 5px rgba(16, 185, 129, 0.5);
        }

        /* Card styling */
        .sleek-card {
            background-color: #1e293b;
            border-radius: 10px;
            padding: 1.5rem;
            margin-bottom: 1.5rem;
            border: 1px solid rgba(148, 163, 184, 0.1);
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
        }

        /* Category pills */
        .category-pill {
            display: inline-block;
            padding: 0.25rem 0.75rem;
            background-color: #1e293b;
            color: #94a3b8;
            border-radius: 9999px;
            font-size: 0.75rem;
            margin-right: 0.5rem;
            margin-bottom: 0.5rem;
            border: 1px solid #334155;
            transition: all 0.3s ease;
        }

        .category-pill:hover {
            background-color: #334155;
            color: #e2e8f0;
            cursor: pointer;
            border-color: #475569;
        }

        /* Chat container */
        .chat-container {
            background-color: #1e293b;
            border-radius: 12px;
            padding: 1.5rem;
            margin-top: 1.5rem;
            height: 500px;
            overflow-y: auto;
            border: 1px solid rgba(148, 163, 184, 0.1);
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
        }

        /* Message styling */
        .message-container {
            display: flex;
            margin-bottom: 1.5rem;
            gap: 12px;
        }

        .user-container {
            justify-content: flex-end;
        }

        .assistant-container {
            justify-content: flex-start;
        }

        .avatar {
            width: 36px;
            height: 36px;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            font-weight: bold;
            font-size: 0.875rem;
            flex-shrink: 0;
        }

        .user-avatar {
            background-color: #7c3aed;
            color: white;
        }

        .assistant-avatar {
            background-color: #10b981;
            color: white;
        }

        .message-box {
            padding: 1rem;
            border-radius: 12px;
            max-width: 90%;
            min-width: 100px;
        }

        .user-message {
            background-color: #7c3aed;
            color: white;
            border-top-right-radius: 0;
        }

        .assistant-message {
            background-color: #334155;
            color: white;
            border-top-left-radius: 0;
        }

        /* Typing indicator animation */
        .typing-indicator {
            display: flex;
            align-items: center;
            padding: 1rem;
            background-color: #334155;
            border-radius: 12px;
            border-top-left-radius: 0;
            min-width: 60px;
        }

        .typing-dot {
            display: block;
            width: 8px;
            height: 8px;
            border-radius: 50%;
            background-color: #94a3b8;
            margin-right: 5px;
            animation: bounce 1.3s linear infinite;
        }

        .typing-dot:nth-child(2) {
            animation-delay: 0.15s;
        }

        .typing-dot:nth-child(3) {
            animation-delay: 0.3s;
            margin-right: 0;
        }

        @keyframes bounce {
            0%, 60%, 100% {
                transform: translateY(0);
            }
            30% {
                transform: translateY(-5px);
            }
        }

        /* Button and input customization */
        .stButton>button {
            background-color: #10b981;
            color: white;
            border: none;
            padding: 0.5rem 1rem;
            border-radius: 8px;
            font-weight: 500;
            transition: all 0.3s ease;
        }

        .stButton>button:hover {
            background-color: #059669;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
        }

        /* Scrollbar styling */
        .chat-container::-webkit-scrollbar {
            width: 6px;
        }

        .chat-container::-webkit-scrollbar-track {
            background: #1e293b;
        }

        .chat-container::-webkit-scrollbar-thumb {
            background-color: #475569;
            border-radius: 20px;
            border: 2px solid #1e293b;
        }

        .chat-container::-webkit-scrollbar-thumb:hover {
            background: #64748b;
        }

        /* Radio button styling for tabs */
        div.row-widget.stRadio > div {
            flex-direction: row;
            gap: 10px;
        }

        div.row-widget.stRadio > div > label {
            background-color: #1e293b;
            border-radius: 8px;
            padding: 6px 16px;
            display: flex;
            align-items: center;
            justify-content: center;
            color: #94a3b8;
            font-weight: 500;
            border: 1px solid #334155;
            transition: all 0.2s ease;
        }

        div.row-widget.stRadio > div > label:hover {
            background-color: #334155;
            color: #e2e8f0;
            cursor: pointer;
        }

        div.row-widget.stRadio > div[role="radiogroup"] input[type="radio"]:checked + label {
            background-color: #10b981;
            color: white;
            border-color: #059669;
        }

        /* Code block styling */
        code {
            background-color: #0f172a;
            padding: 0.2em 0.4em;
            border-radius: 4px;
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.85em;
        }

        pre {
            background-color: #0f172a;
            padding: 1rem;
            border-radius: 8px;
            overflow-x: auto;
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.85em;
            border: 1px solid #334155;
        }

        /* Hide Streamlit branding */
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        </style>
        """, unsafe_allow_html=True)
    # Add custom CSS for modern UI - Kept original CSS from your code


    st.markdown("""
    <style>
    html, body, [data-testid="stAppViewContainer"], .main {
        background: linear-gradient(to bottom, rgb(3, 7, 18), rgb(17, 24, 39)) !important;
        color: #10b981 !important;
    }
    body {
        background-color: #000 !important;
        color: #10b981 !important;
        font-family: 'Inter', sans-serif;
    }
    .main-header {
        padding: 1.5rem;
        background: linear-gradient(to bottom, rgb(3, 7, 18), rgb(17, 24, 39));
        border-radius: 12px;
        margin-bottom: 1.5rem;
        border: 1px solid #10b98144;
        box-shadow: 0 4px 6px -1px #10b98122, 0 2px 4px -1px #10b98122;
    }
    .status-indicator {
        width: 12px;
        height: 12px;
        background-color: #10b981;
        border-radius: 50%;
        display: inline-block;
        margin-right: 8px;
        box-shadow: 0 0 0 #10b98177;
        animation: pulse 2s infinite;
    }
    @keyframes pulse {
        0% { box-shadow: 0 0 0 0 #10b98166; }
        70% { box-shadow: 0 0 0 10px #10b98100; }
        100% { box-shadow: 0 0 0 0 #10b98100; }
    }
    .glow-text {
        color: #10b981;
        font-weight: 500;
        text-shadow: 0 0 8px #10b98188;
    }
    .sleek-card {
        background-color: #101010;
        border-radius: 10px;
        padding: 1.5rem;
        margin-bottom: 1.5rem;
        border: 1px solid #10b98144;
        box-shadow: 0 4px 6px -1px #10b98122, 0 2px 4px -1px #10b98122;
    }
    .category-pill {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        background-color: #101010;
        color: #10b981;
        border-radius: 9999px;
        font-size: 0.75rem;
        margin-right: 0.5rem;
        margin-bottom: 0.5rem;
        border: 1px solid #10b98177;
        transition: all 0.3s ease;
    }
    .category-pill:hover {
        background-color: #10b98122;
        color: #fff;
        cursor: pointer;
        border-color: #10b981;
    }
    .chat-container {
        background-color: #101010;
        border-radius: 12px;
        padding: 1.5rem;
        margin-top: 1.5rem;
        height: 500px;
        overflow-y: auto;
        border: 1px solid #10b98144;
        box-shadow: 0 4px 6px -1px #10b98122, 0 2px 4px -1px #10b98122;
    }
    .message-container {
        display: flex;
        margin-bottom: 1.5rem;
        gap: 12px;
    }
    .user-container { justify-content: flex-end; }
    .assistant-container { justify-content: flex-start; }
    .avatar {
        width: 36px;
        height: 36px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: bold;
        font-size: 0.875rem;
        flex-shrink: 0;
        background: #10b981;
        color: #000;
    }
    .user-avatar { background-color: #10b981; color: #000; }
    .assistant-avatar { background-color: #222; color: #10b981; }
    .message-box {
        padding: 1rem;
        border-radius: 12px;
        max-width: 90%;
        min-width: 100px;
        background: #101010;
        color: white;
        border: 1px solid #10b98144;
    }
    .user-message {
        background-color: #10b981;
        color: #000;
        border-top-right-radius: 0;
    }
    .assistant-message {
        background-color: #101010;
        color: white;
        border-top-left-radius: 0;
        border: 1px solid #10b98144;
    }
    .typing-indicator {
        display: flex;
        align-items: center;
        padding: 1rem;
        background-color: #101010;
        border-radius: 12px;
        border-top-left-radius: 0;
        min-width: 60px;
        border: 1px solid #10b98144;
    }
    .typing-dot {
        display: block;
        width: 8px;
        height: 8px;
        border-radius: 50%;
        background-color: #10b981;
        margin-right: 5px;
        animation: bounce 1.3s linear infinite;
    }
    .typing-dot:nth-child(2) { animation-delay: 0.15s; }
    .typing-dot:nth-child(3) { animation-delay: 0.3s; margin-right: 0; }
    @keyframes bounce {
        0%, 60%, 100% { transform: translateY(0); }
        30% { transform: translateY(-5px); }
    }
    .stButton>button {
        background-color: #10b981;
        color: #000;
        border: none;
        padding: 0.5rem 1rem;
        border-radius: 8px;
        font-weight: 500;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #13d49d;
        color: #000;
        box-shadow: 0 4px 6px -1px #10b98122, 0 2px 4px -1px #10b98122;
    }
    .chat-container::-webkit-scrollbar { width: 6px; }
    .chat-container::-webkit-scrollbar-track { background: #101010; }
    .chat-container::-webkit-scrollbar-thumb {
        background-color: #10b98144;
        border-radius: 20px;
        border: 2px solid #101010;
    }
    .chat-container::-webkit-scrollbar-thumb:hover { background: #10b981; }
    div.row-widget.stRadio > div > label {
        background-color: #101010;
        border-radius: 8px;
        padding: 6px 16px;
        color: #10b981;
        font-weight: 500;
        border: 1px solid #10b98144;
        transition: all 0.2s ease;
    }
    div.row-widget.stRadio > div > label:hover {
        background-color: #10b98122;
        color: #fff;
        cursor: pointer;
    }
    div.row-widget.stRadio > div[role="radiogroup"] input[type="radio"]:checked + label {
        background-color: #10b981;
        color: #000;
        border-color: #13d49d;
    }
    code, pre {
        background-color: #000 !important;
        color: #10b981 !important;
        border-radius: 4px;
        font-family: 'JetBrains Mono', monospace;
        font-size: 0.95em;
    }
    pre {
        padding: 1rem;
        border-radius: 8px;
        overflow-x: auto;
        border: 1px solid #10b98144;
    }
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    </style>
""", unsafe_allow_html=True)

    # Initialize session state
    if 'assistant' not in st.session_state:
        st.session_state.assistant = None
    
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    
    if 'thinking' not in st.session_state:
        st.session_state.thinking = False
        
    if 'current_tab' not in st.session_state:
        st.session_state.current_tab = "Chat"
        
    # Header with cyber security themed logo
    st.markdown("""
        <div class="main-header">
            <h1 style="font-size: 2.75rem; color: #e2e8f0; margin-bottom: 0.5rem;">GamkersGPT AI</h1>
            <p style="color: #a0aec0; font-size: 1.1rem;">
                AI-powered cybersecurity training assistant for ethical hacking education by gamkers
            </p>
            <div style="margin-top: 1rem; display: flex; align-items: center;">
                <div class="status-indicator"></div>
                <span class="glow-text">Active & Secure</span>
            </div>
        </div>
    """, unsafe_allow_html=True)

    try:
        # Get API key from secrets.toml or environment variable
        groq_api_key=st.secrets["groq_api_key"]
                
        # Initialize the assistant
        st.session_state.assistant = CyberSecurityAssistant(groq_api_key=groq_api_key)
        st.success("GamkersGPT initialized successfully!")
                
    except Exception as e:
        st.error(f"Error initializing assistant: {e}")

    def chat():    
        # Main content - Tabs system
        tabs = ["Chat", "CVE Database", "Security News", "Tool Commands", "Google Dorking"]
        st.session_state.current_tab = st.radio("Select Feature", tabs, horizontal=True)
        
        # Chat Tab
        if st.session_state.current_tab == "Chat":
            # Feature cards
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("""
                    <div class="sleek-card">
                        <h3 style="margin: 0 0 0.75rem 0; color: #e2e8f0; display: flex; align-items: center;">
                            <svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#f59e0b" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="margin-right: 0.75rem;">
                                <path d="M10.29 3.86L1.82 18a1 1 0 0 0 .86 1.5h18.64a1 1 0 0 0 .86-1.5L13.71 3.86a1 1 0 0 0-1.72 0z"></path>
                                <line x1="12" y1="9" x2="12" y2="13"></line>
                                <line x1="12" y1="17" x2="12.01" y2="17"></line>
                            </svg>
                            Disclaimer: Ethical Use Only
                        </h3>
                        <p style="color: #a0aec0; margin-bottom: 0;">All cybersecurity content provided here is strictly for educational and ethical purposes. We do not support or encourage any unauthorized or illegal access to systems or networks. Always obtain proper permission before conducting any form of security testing.</p>
                    </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown("""
                    <div class="sleek-card">
                        <h3 style="margin: 0 0 0.75rem 0; color: #e2e8f0; display: flex; align-items: center;">
                            <svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#10b981" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="margin-right: 0.75rem;">
                                <path d="M12 22s8-4 8-10V5l-8-3-8 3v7c0 6 8 10 8 10z"></path>
                            </svg>
                            Security Best Practices
                        </h3>
                        <p style="color: #a0aec0; margin-bottom: 0;">Learn defensive security measures, incident response protocols, and security best practices to protect systems and networks from cyber threats.</p>
                    </div>
                """, unsafe_allow_html=True)
                
            # Chat container
            # st.markdown('<div class="chat-container">', unsafe_allow_html=True)
            
            # Function to display messages
            def display_messages():
                for message in st.session_state.messages:
                    if message["role"] == "user":
                        st.markdown(
                            f"""
                            <div class="message-container user-container">
                                <div class="avatar user-avatar">U</div>
                                <div class="message-box user-message">{message['content']}</div>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )
                    elif message["role"] == "assistant":
                        st.markdown(
                            f"""
                            <div class="message-container assistant-container">
                                <div class="avatar assistant-avatar">C</div>
                                <div class="message-box assistant-message">{message['content'].replace("**","").replace("=====================================","")}</div>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )
                
                # Display thinking animation if assistant is processing
                if st.session_state.thinking:
                    st.markdown(
                        """
                        <div class="message-container assistant-container">
                            <div class="avatar assistant-avatar">C</div>
                            <div class="typing-indicator">
                                <div class="typing-dot"></div>
                                <div class="typing-dot"></div>
                                <div class="typing-dot"></div>
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
        
            # Display previous messages
            display_messages()
            
            # Initial hint if no messages
            if not st.session_state.messages:
                st.markdown("""
                    <div style="text-align: center; padding: 32px 16px; color: #6b7280;">
                        <svg xmlns="http://www.w3.org/2000/svg" width="48" height="48" viewBox="0 0 24 24" fill="none" stroke="#10b981" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round" style="margin: 0 auto 16px;">
                            <path d="M21 15a2 2 0 0 1-2 2H7l-4 4V5a2 2 0 0 1 2-2h14a2 2 0 0 1 2 2z"></path>
                        </svg>
                        <p>No messages yet. Initialize the assistant and start the conversation!</p>
                    </div>
                """, unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            # User input and processing
            if prompt := st.chat_input("Ask GamkersGPT...", key="user_input"):
                if st.session_state.assistant is None:
                    st.warning("Please initialize the assistant first using the button in the sidebar.")
                else:
                    # Add user message to chat
                    st.session_state.messages.append({"role": "user", "content": prompt})
                    
                    # Set thinking state to true and rerun to show the animation
                    st.session_state.thinking = True
                    st.rerun()
            
            # This part only runs after the rerun if thinking is True
            if st.session_state.thinking and st.session_state.assistant is not None:
                try:
                    # Get assistant response
                    response = st.session_state.assistant.query(st.session_state.messages[-1]["content"])
                    
                    # Add assistant response
                    st.session_state.messages.append({"role": "assistant", "content": response})
                    
                    # Stop thinking, rerun to show updated chat
                    st.session_state.thinking = False
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Error processing your request: {str(e)}")
                    st.session_state.thinking = False
                    st.rerun()
        
        # CVE Database Tab
        elif st.session_state.current_tab == "CVE Database":
            st.markdown("""
                <div class="sleek-card">
                    <h3 style="margin: 0 0 0.75rem 0; color: #e2e8f0; display: flex; align-items: center;">
                        <svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#10b981" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="margin-right: 0.75rem;">
                            <circle cx="11" cy="11" r="8"></circle>
                            <line x1="21" y1="21" x2="16.65" y2="16.65"></line>
                        </svg>
                        CVE Database Search
                    </h3>
                    <p style="color: #a0aec0; margin-bottom: 1rem;">Search for Common Vulnerabilities and Exposures (CVE) by ID or keyword. Get detailed information about security vulnerabilities, including severity, impact, and remediation.</p>
                </div>
            """, unsafe_allow_html=True)
            
            # CVE search form
            search_col1, search_col2 = st.columns([3, 1])
            
            with search_col1:
                cve_query = st.text_input("Enter CVE ID or keyword", 
                                        placeholder="Example: CVE-2021-44228 or log4j")
                
            with search_col2:
                cve_limit = st.selectbox("Result limit", [5, 10, 15, 20], index=0)
                search_button = st.button("Search CVE", key="search_cve_button")
            
            # Search process
            if search_button and cve_query:
                if st.session_state.assistant is None:
                    st.warning("Please initialize the assistant first using the button in the sidebar.")
                else:
                    with st.spinner("Searching CVE database..."):
                        cve_results = st.session_state.assistant.search_cve(cve_query, cve_limit)
                        
                        if isinstance(cve_results, str):
                            st.warning(cve_results)
                        else:
                            # Display results in a nice format
                            for result in cve_results:
                                with st.expander(f"{result['cve_id']} - Severity: {result['severity']} ({result['base_score']})"):
                                    st.markdown(f"### {result['cve_id']}")
                                    
                                    # Create two columns for severity and dates
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        st.markdown(f"**Severity**: {result['severity']}")
                                        st.markdown(f"**CVSS Score**: {result['base_score']}")
                                    
                                    with col2:
                                        st.markdown(f"**Published**: {result['published']}")
                                        st.markdown(f"**Last Modified**: {result['last_modified']}")
                                    
                                    st.markdown("### Description")
                                    st.markdown(result['description'])
                                    
                                    # References
                                    if result['references']:
                                        st.markdown("### References")
                                        for ref in result['references'][:5]:  # Limit to first 5 references
                                            st.markdown(f"- [{ref.get('url', 'Link')}]({ref.get('url', '#')})")
        
        # Security News Tab
        elif st.session_state.current_tab == "Security News":
            st.markdown("""
                <div class="sleek-card">
                    <h3 style="margin: 0 0 0.75rem 0; color: #e2e8f0; display: flex; align-items: center;">
                        <svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#10b981" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="margin-right: 0.75rem;">
                            <path d="M19 20H5a2 2 0 0 1-2-2V6a2 2 0 0 1 2-2h10l6 6v8a2 2 0 0 1-2 2z"></path>
                            <line x1="12" y1="18" x2="12" y2="12"></line>
                            <line x1="9" y1="15" x2="15" y2="15"></line>
                        </svg>
                        Latest Cybersecurity News
                    </h3>
                    <p style="color: #a0aec0; margin-bottom: 1rem;">Stay updated with the latest cybersecurity news from The Hacker News, summarized by AI for quick insights into current threats, vulnerabilities, and industry developments.</p>
                </div>
            """, unsafe_allow_html=True)
            
            # News fetch options
            news_col1, news_col2 = st.columns([3, 1])
            
            with news_col2:
                news_limit = st.selectbox("Article limit", [3, 5, 7, 10], index=1)
                news_button = st.button("Fetch News", key="fetch_news_button")
            
            # News fetch process
            if news_button:
                if st.session_state.assistant is None:
                    st.warning("Please initialize the assistant first using the button in the sidebar.")
                else:
                    with st.spinner("Fetching and summarizing latest cybersecurity news..."):
                        news_results = st.session_state.assistant.get_hacker_news(news_limit)
                        
                        if isinstance(news_results, str):
                            st.warning(news_results)
                        else:
                            # Display news in a nice format
                            for i, news in enumerate(news_results):
                                with st.container():
                                    st.markdown(f"""
                                    <div class="sleek-card" style="margin-bottom: 1rem;">
                                        <h4 style="margin: 0 0 0.5rem 0; color: #e2e8f0;">{news['title']}</h4>
                                        <p style="color: #a0aec0; font-size: 0.8rem; margin-bottom: 0.75rem;">Published: {news['published']}</p>
                                        <p style="color: #e2e8f0; margin-bottom: 1rem;">{news['summary']}</p>
                                        <a href="{news['link']}" target="_blank" style="color: #10b981; text-decoration: none; font-weight: 500;">Read full article →</a>
                                    </div>
                                    """, unsafe_allow_html=True)
        
        # Tool Commands Tab
        elif st.session_state.current_tab == "Tool Commands":
            st.markdown("""
                <div class="sleek-card">
                    <h3 style="margin: 0 0 0.75rem 0; color: #e2e8f0; display: flex; align-items: center;">
                        <svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#10b981" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="margin-right: 0.75rem;">
                            <path d="M14.7 6.3a1 1 0 0 0 0 1.4l1.6 1.6a1 1 0 0 0 1.4 0l3.77-3.77a6 6 0 0 1-7.94 7.94l-6.91 6.91a2.12 2.12 0 0 1-3-3l6.91-6.91a6 6 0 0 1 7.94-7.94l-3.76 3.76z"></path>
                        </svg>
                        Security Tool Commands
                    </h3>
                    <p style="color: #a0aec0; margin-bottom: 1rem;">Generate practical command examples for common security tools. Learn how to use tools like Nmap, Metasploit, Wireshark, and more for your security assessments.</p>
                </div>
            """, unsafe_allow_html=True)
            
            # Tool commands form
            tool_col1, tool_col2 = st.columns(2)
            
            with tool_col1:
                tool_name = st.text_input("Enter security tool name", placeholder="Example: nmap, metasploit, burpsuite")
                
            with tool_col2:
                task_desc = st.text_input("What do you want to achieve with this tool?", 
                                        placeholder="Example: network scanning, discovery of web vulnerabilities")
            
            tool_button = st.button("Generate Commands", key="gen_tool_button")
            
            # Tool command generation process
            if tool_button and tool_name and task_desc:
                if st.session_state.assistant is None:
                    st.warning("Please initialize the assistant first using the button in the sidebar.")
                else:
                    with st.spinner(f"Generating {tool_name} commands for {task_desc}..."):
                        tool_results = st.session_state.assistant.generate_tool_commands(tool_name, task_desc)
                        # st.write(tool_results)
                        st.markdown(f"""
                            <div class="sleek-card" style="background-color: #0f172a; border: 1px solid #1e293b;">
                                <h4 style="margin: 0 0 1rem 0; color: #e2e8f0; display: flex; align-items: center;">
                                    <svg xmlns="http://www.w3.org/2000/svg" width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="#10b981" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="margin-right: 0.75rem;">
                                        <polyline points="9 11 12 14 22 4"></polyline>
                                        <path d="M21 12v7a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h11"></path>
                                    </svg>
                                    Generated Commands: {tool_name}
                                </h4>
                                <div style="color: #e2e8f0; font-family: monospace; background-color: #1e293b; padding: 1rem; border-radius: 8px; white-space: pre-wrap;">
                                    {tool_results}
                                </div>
                            </div>
                        """, unsafe_allow_html=True)
        
        # Google Dorking Tab
        elif st.session_state.current_tab == "Google Dorking":
            st.markdown("""
                <div class="sleek-card">
                    <h3 style="margin: 0 0 0.75rem 0; color: #e2e8f0; display: flex; align-items: center;">
                        <svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="#10b981" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="margin-right: 0.75rem;">
                            <circle cx="11" cy="11" r="8"></circle>
                            <line x1="21" y1="21" x2="16.65" y2="16.65"></line>
                        </svg>
                        Google Dorking Generator
                    </h3>
                    <p style="color: #a0aec0; margin-bottom: 1rem;">Generate Google dorking queries for finding specific information on websites. Learn advanced search operators to identify potential security vulnerabilities and improve targeted reconnaissance.</p>
                </div>
            """, unsafe_allow_html=True)
            
            # Google dork form
            dork_col1, dork_col2 = st.columns(2)
            
            with dork_col1:
                target_type = st.text_input("What are you looking for?", 
                                        placeholder="Example: exposed databases, login pages, config files")
                
            with dork_col2:
                dork_objective = st.text_input("What's your security objective?", 
                                            placeholder="Example: vulnerability assessment, information gathering")
            
            dork_button = st.button("Generate Google Dorks", key="gen_dork_button")
            
            # Dork generation process
            if dork_button and target_type and dork_objective:
                if st.session_state.assistant is None:
                    st.warning("Please initialize the assistant first using the button in the sidebar.")
                else:
                    with st.spinner(f"Generating Google dorks for finding {target_type}..."):
                        dork_results = st.session_state.assistant.generate_google_dorks(target_type, dork_objective)
                        
                        st.write(dork_results)

    try:
        with open("style.css") as f:
            st.markdown('<style>' + f.read() + '</style>', unsafe_allow_html=True)
    except Exception:
        pass
    with st.sidebar:
        tabs = on_hover_tabs(
        tabName=['Encryption', 'chat', 'Security Assessment', 'Training', 'Malware Analysis'],
        iconName=['lock', 'comments', 'shield-alt', 'home', 'code'],
        styles={
            'navtab': {
                'background-color': 'black',
                'color': 'White', 
                'font-size': '16px',
                'transition': '.3s',
                'white-space': 'nowrap',
                'text-transform': 'uppercase'
            },
            'tabOptionsStyle': {
                ':hover': {'color': '#1A1A1A', 'background-color': 'black'}
            },
            'iconStyle': {
                'position': 'fixed',
                'left': '7.5px',
                'text-align': 'left'
            },
            'tabStyle': {
                'list-style-type': 'none',
                'margin-bottom': '30px',
                'padding-left': '30px'
            }
        }
    )

    if tabs == 'chat':
        chat()
    elif tabs == 'Encryption':
        if True:
            crypto_op = st.radio("Operation", ["Encode/Decode", "Hash Analysis", "Encrypt/Decrypt"])
            
            if crypto_op == "Encode/Decode":
                # Encoding/Decoding tool
                encode_decode_op = st.radio("Action", ["Encode", "Decode"], horizontal=True)
                format_type = st.selectbox("Format", ["base64", "hex", "url", "binary"])
                input_data = st.text_area("Input Data", height=100)
                
                if st.button("Process", key="encode_decode_btn"):
                    if st.session_state.assistant is None:
                        st.warning("Please initialize assistant first")
                    elif not input_data:
                        st.warning("Please enter data to process")
                    else:
                        result = st.session_state.assistant.decode_or_encode_data(
                            input_data, 
                            encode_decode_op.lower(), 
                            format_type
                        )
                        
                        if isinstance(result, dict):
                            st.code(result["result"], language="text")
                            st.success(result["explanation"])
                        else:
                            st.error(result)
            
            elif crypto_op == "Hash Analysis":
                # Hash Analysis Tool
                hash_input = st.text_area("Enter Hash to Analyze", height=80)
                
                if st.button("Analyze Hash", key="analyze_hash_btn"):
                    if st.session_state.assistant is None:
                        st.warning("Please initialize assistant first")
                    elif not hash_input:
                        st.warning("Please enter a hash to analyze")
                    else:
                        result = st.session_state.assistant.hash_analyzer(hash_input)
                        
                        if isinstance(result, dict):
                            st.markdown(f"""
                            **Hash Type:** {result["hash_type"]}\n
                            **Length:** {result["length"]}\n
                            **Security Strength:** {result["strength"]}\n
                            **Valid Format:** {result["valid_format"]}\n
                            **Additional Info:** {result["additional_info"]}
                            """)
                        else:
                            st.error(result)
            
            elif crypto_op == "Encrypt/Decrypt":
                # Encryption/Decryption Tool
                crypto_action = st.radio("Action", ["Encrypt", "Decrypt"], horizontal=True)
                algorithm = st.selectbox("Algorithm", ["aes", "des", "chacha20"])
                crypto_input = st.text_area("Input Data", height=80)
                
                key_col, iv_col = st.columns(2)
                with key_col:
                    key_input = st.text_input("Key (Base64, leave empty to generate)")
                with iv_col:
                    iv_input = st.text_input("IV/Nonce (Base64, leave empty to generate)")
                
                if st.button("Process", key="crypto_process_btn"):
                    if st.session_state.assistant is None:
                        st.warning("Please initialize assistant first")
                    elif not crypto_input:
                        st.warning("Please enter data to process")
                    else:
                        result = st.session_state.assistant.decrypt_or_encrypt_data(
                            crypto_input,
                            crypto_action.lower(),
                            algorithm,
                            key_input if key_input else None,
                            iv_input if iv_input else None
                        )
                        
                        if isinstance(result, dict):
                            st.code(result["result"], language="text")
                            
                            if crypto_action.lower() == "encrypt":
                                st.info("Save these values for decryption:")
                                st.code(f"Key: {result['key']}\nIV: {result['iv']}", language="text")
                        else:
                            st.error(result)

    elif tabs == 'Security Assessment':
        if True:
            assessment_type = st.selectbox(
                "Assessment Type",
                ["Vulnerability Assessment", "Code Security Analysis", "Network Traffic Analysis"]
            )
            
            if assessment_type == "Vulnerability Assessment":
                system_desc = st.text_area("Describe the system to assess", height=100)
                
                if st.button("Generate Assessment", key="vuln_assess_btn"):
                    if st.session_state.assistant is None:
                        st.warning("Please initialize assistant first")
                    elif not system_desc:
                        st.warning("Please describe the system to assess")
                    else:
                        with st.spinner("Generating vulnerability assessment..."):
                            result = st.session_state.assistant.vulnerability_assessment(system_desc)
                            st.markdown(result)
            
            elif assessment_type == "Code Security Analysis":
                code_lang = st.selectbox("Language", ["python", "javascript", "php", "java", "html", "css"])
                code_snippet = st.text_area("Paste code to analyze", height=150)
                
                if st.button("Analyze Code", key="code_analysis_btn"):
                    if st.session_state.assistant is None:
                        st.warning("Please initialize assistant first")
                    elif not code_snippet:
                        st.warning("Please enter code to analyze")
                    else:
                        with st.spinner("Analyzing code security..."):
                            result = st.session_state.assistant.analyze_code_security(code_snippet, code_lang)
                            st.markdown(result)
            
            elif assessment_type == "Network Traffic Analysis":
                traffic_type = st.selectbox("Traffic Type", ["http", "dns", "tcp", "ssl", "smb"])
                traffic_desc = st.text_area("Describe traffic pattern or paste sample data", height=100)
                
                if st.button("Analyze Traffic", key="traffic_analysis_btn"):
                    if st.session_state.assistant is None:
                        st.warning("Please initialize assistant first")
                    else:
                        with st.spinner("Analyzing network traffic patterns..."):
                            result = st.session_state.assistant.analyze_network_traffic(
                                traffic_desc if traffic_desc else None,
                                traffic_type
                            )
                            st.markdown(result)
    elif tabs=='Training':
        if True:
            training_type = st.selectbox(
                "Material Type",
                ["Security Concept Explanation", "CTF Challenge", "Incident Response Plan", "Security Policy"]
            )
            
            if training_type == "Security Concept Explanation":
                concept = st.text_input("Enter security concept", placeholder="Example: XSS, CSRF, JWT")
                
                if st.button("Generate Explanation", key="concept_explain_btn"):
                    if st.session_state.assistant is None:
                        st.warning("Please initialize assistant first")
                    elif not concept:
                        st.warning("Please enter a security concept")
                    else:
                        with st.spinner(f"Generating explanation for {concept}..."):
                            # Get the text explanation from your assistant
                            result = st.session_state.assistant.explain_security_concept(concept)
                            
                            # Display the explanation
                            st.markdown("## 📚 Concept Explanation")
                            st.markdown(result)
                            
                            # Add resource search section
                            st.markdown("## 🔍 Additional Resources")
                            
                            # Create columns for different resource types
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                st.markdown("### 📄 PDF Documents")
                                with st.spinner("Searching for PDFs..."):
                                    pdf_queries = [
                                        f"{concept} security guide",
                                        f"{concept} tutorial PDF",
                                        f"{concept} whitepaper",
                                        f"{concept} OWASP guide"
                                    ]
                                    
                                    pdf_found = False
                                    for i, query in enumerate(pdf_queries):
                                        try:
                                            pdf_result = search_pdf(query)
                                            if "PDF found:" in pdf_result:
                                                pdf_link = pdf_result.replace("PDF found: ", "")
                                                st.markdown(f"📄 [{query.title()}]({pdf_link})")
                                                pdf_found = True
                                                if i >= 2:  # Limit to 3 PDFs
                                                    break
                                        except:
                                            continue
                                    
                                    if not pdf_found:
                                        st.info("No PDF documents found")
                            
                            with col2:
                                st.markdown("### 🎯 PowerPoint Presentations")
                                with st.spinner("Searching for presentations..."):
                                    ppt_queries = [
                                        f"{concept} security presentation",
                                        f"{concept} slides tutorial",
                                        f"{concept} PowerPoint cybersecurity"
                                    ]
                                    
                                    ppt_found = False
                                    for i, query in enumerate(ppt_queries):
                                        try:
                                            ppt_result = search_ppt(query)
                                            if "PPT found:" in ppt_result:
                                                ppt_link = ppt_result.replace("PPT found: ", "")
                                                st.markdown(f"📊 [{query.title()}]({ppt_link})")
                                                ppt_found = True
                                                if i >= 2:  # Limit to 3 PPTs
                                                    break
                                        except:
                                            continue
                                    
                                    if not ppt_found:
                                        st.info("No presentations found")
                            
                            # Research Papers Section
                            st.markdown("### 🔬 Research Papers")
                            with st.spinner("Searching for research papers..."):
                                research_queries = [
                                    f"{concept} security research paper",
                                    f"{concept} academic paper cybersecurity",
                                    f"{concept} vulnerability research"
                                ]
                                
                                research_found = False
                                for i, query in enumerate(research_queries):
                                    try:
                                        research_result = search_pdf(query)
                                        if "PDF found:" in research_result:
                                            research_link = research_result.replace("PDF found: ", "")
                                            st.markdown(f"📋 [{query.title()}]({research_link})")
                                            research_found = True
                                            if i >= 2:  # Limit to 3 research papers
                                                break
                                    except:
                                        continue
                                
                                if not research_found:
                                    st.info("No research papers found")
                            
                            # YouTube Videos Section - Embedded
                            st.markdown("### 📺 Video Tutorials")
                            with st.spinner("Searching and loading videos..."):
                                video_queries = [
                                    f"{concept}",
                                    f"{concept} full course",
                                    f"{concept} attack demonstration",
                                    f"how to prevent {concept} attacks"
                                ]
                                
                                videos_found = 0
                                video_cols = st.columns(2)  # 2 columns for videos
                                
                                for i, query in enumerate(video_queries):
                                    try:
                                        video_result = search_youtube(query)
            
                                        if video_result:
                                            with video_cols[videos_found % 2]:
                                                st.markdown(f"**{query.title()}**")
                                                st.video(video_result)
                                            
                                            videos_found += 1
                                            if videos_found >= 4:  # Limit to 4 videos
                                                break
                                    except:
                                        continue
                                
                                if videos_found == 0:
                                    st.info("No videos found for this concept")

                            st.markdown("---")
                            st.markdown("### 📝 Security Blogs & Articles")
                            blog_col1, blog_col2 = st.columns(2)
                            
                            with blog_col1:
                                st.markdown("#### Technical Blogs")
                                with st.spinner("Searching for technical blogs..."):
                                    blog_queries = [
                                        f"{concept} security blog",
                                        f"{concept} cybersecurity article",
                                        f"{concept} penetration testing blog",
                                        f"how to exploit {concept}"
                                    ]
                                    
                                    blogs_found = False
                                    for query in blog_queries:
                                        try:
                                            blog_result = search_blogs(query)
                                            if "Blog found:" in blog_result:
                                                blog_link = blog_result.replace("Blog found: ", "")
                                                st.markdown(f"📝 [{query.title()}]({blog_link})")
                                                blogs_found = True
                                        except:
                                            continue
                                    
                                    if not blogs_found:
                                        st.info("No technical blogs found")
                            
                            with blog_col2:
                                st.markdown("#### Security Research Articles")
                                with st.spinner("Searching for research articles..."):
                                    research_blog_queries = [
                                        f"{concept} POC",
                                        f"{concept} vulnerability analysis",
                                        f"{concept} security analysis blog",
                                        f"{concept} attack analysis"
                                    ]
                                    
                                    research_blogs_found = False
                                    for query in research_blog_queries:
                                        try:
                                            research_blog_result = search_poc_reports(query)
                                            if "PoC/Report found:" in research_blog_result:
                                                research_blog_link = research_blog_result.replace("PoC/Report found: ", "")
                                                st.markdown(f"🔬 [{query.title()}]({research_blog_link})")
                                                research_blogs_found = True
                                        except:
                                            continue
                                    
                                    if not research_blogs_found:
                                        st.info("No research articles found")
                            
                            # Practice Labs Section
                            st.markdown("---")
                            st.markdown("### 🧪 Practice Labs & Exercises")
                            with st.spinner("Searching for hands-on labs..."):
                                lab_queries = [
                                    f"{concept} lab exercise",
                                    f"{concept} hands-on tutorial",
                                    f"{concept} practical demonstration",
                                    f"{concept} CTF challenge"
                                ]
                                
                                labs_found = False
                                for query in lab_queries:
                                    try:
                                        lab_result = search_pdf(query)
                                        if "PDF found:" in lab_result:
                                            lab_link = lab_result.replace("PDF found: ", "")
                                            st.markdown(f"🧪 [{query.title()}]({lab_link})")
                                            labs_found = True
                                    except:
                                        continue
                                
                                if not labs_found:
                                    st.info("No practice labs found")
                            
                            # Assessment Materials Section
                            st.markdown("### 📋 Assessment Materials")
                            with st.spinner("Searching for assessment materials..."):
                                assessment_queries = [
                                    f"{concept} quiz questions",
                                    f"{concept} exam preparation",
                                    f"{concept} practice test",
                                    f"{concept} assessment PDF"
                                ]
                                
                                assessments_found = False
                                for query in assessment_queries:
                                    try:
                                        assessment_result = search_pdf(query)
                                        if "PDF found:" in assessment_result:
                                            assessment_link = assessment_result.replace("PDF found: ", "")
                                            st.markdown(f"📋 [{query.title()}]({assessment_link})")
                                            assessments_found = True
                                    except:
                                        continue
                                
                                if not assessments_found:
                                    st.info("No assessment materials found")
                            
                            # Additional Resources Section
                            st.markdown("---")
                            st.markdown("### 🌐 Additional Learning Resources")
                            
                            additional_col1, additional_col2 = st.columns(2)
                            
                            with additional_col1:
                                st.markdown("#### Security Tools & Documentation")
                                with st.spinner("Searching for tools and documentation..."):
                                    tool_queries = [
                                        f"{concept} security tools",
                                        f"{concept} detection tools",
                                        f"{concept} prevention tools"
                                    ]
                                    
                                    for query in tool_queries:
                                        try:
                                            tool_result = search_pdf(query)
                                            if "PDF found:" in tool_result:
                                                tool_link = tool_result.replace("PDF found: ", "")
                                                st.markdown(f"🔧 [{query.title()}]({tool_link})")
                                        except:
                                            continue
                            
                            with additional_col2:
                                st.markdown("#### Industry Standards & Best Practices")
                                with st.spinner("Searching for standards and best practices..."):
                                    standard_queries = [
                                        f"{concept} best practices",
                                        f"{concept} security standards",
                                        f"{concept} compliance guide"
                                    ]
                                    
                                    for query in standard_queries:
                                        try:
                                            standard_result = search_pdf(query)
                                            if "PDF found:" in standard_result:
                                                standard_link = standard_result.replace("PDF found: ", "")
                                                st.markdown(f"📜 [{query.title()}]({standard_link})")
                                        except:
                                            continue
                            
                            # Save concept for future reference
                            if 'studied_concepts' not in st.session_state:
                                st.session_state.studied_concepts = []
                            
                            if concept not in st.session_state.studied_concepts:
                                st.session_state.studied_concepts.append(concept)
                                st.success(f"✅ {concept} added to your study history!")
                            
                            # Show study progress
                            if len(st.session_state.studied_concepts) > 0:
                                st.markdown("---")
                                st.markdown("### 📈 Your Study Progress")
                                
                                progress_col1, progress_col2 = st.columns(2)
                                
                                with progress_col1:
                                    st.write(f"**Concepts studied:** {len(st.session_state.studied_concepts)}")
                                    st.write(f"**Current concept:** {concept}")
                                    
                                with progress_col2:
                                    progress_value = min(len(st.session_state.studied_concepts) / 10, 1.0)
                                    st.progress(progress_value)
                                    st.caption(f"Progress: {len(st.session_state.studied_concepts)}/10 concepts mastered")
                                
                                # Display studied concepts
                                if len(st.session_state.studied_concepts) > 1:
                                    st.markdown("**Previously studied concepts:**")
                                    concept_tags = " • ".join(st.session_state.studied_concepts[:-1])  # Exclude current concept
                                    st.write(concept_tags)
                        
            elif training_type == "CTF Challenge":
                ctf_difficulty = st.select_slider(
                    "Difficulty", 
                    options=["easy", "medium", "hard"],
                    value="medium"
                )
                ctf_category = st.selectbox(
                    "Category",
                    ["web", "crypto", "forensics", "pwn", "reverse", "misc"]
                )
                
                if st.button("Generate Challenge", key="ctf_gen_btn"):
                    if st.session_state.assistant is None:
                        st.warning("Please initialize assistant first")
                    else:
                        with st.spinner(f"Generating {ctf_difficulty} {ctf_category} CTF challenge..."):
                            result = st.session_state.assistant.generate_ctf_challenge(
                                ctf_difficulty,
                                ctf_category
                            )
                            st.markdown(result)
            
            elif training_type == "Incident Response Plan":
                ir_incident = st.selectbox(
                    "Incident Type",
                    ["data breach", "ransomware", "DDoS", "insider threat", "phishing"]
                )
                ir_size = st.selectbox("Organization Size", ["small", "medium", "large"])
                
                if st.button("Generate IR Plan", key="ir_gen_btn"):
                    if st.session_state.assistant is None:
                        st.warning("Please initialize assistant first")
                    else:
                        with st.spinner(f"Generating incident response plan..."):
                            result = st.session_state.assistant.generate_incident_response_plan(
                                ir_incident,
                                ir_size
                            )
                            st.markdown(result)
            
            elif training_type == "Security Policy":
                policy_org = st.selectbox(
                    "Organization Type",
                    ["healthcare", "finance", "education", "retail", "government", "technology"]
                )
                policy_focus = st.selectbox(
                    "Policy Focus",
                    ["network", "data", "access control", "incident response", "BYOD", "cloud"]
                )
                
                if st.button("Generate Policy", key="policy_gen_btn"):
                    if st.session_state.assistant is None:
                        st.warning("Please initialize assistant first")
                    else:
                        with st.spinner(f"Generating security policy..."):
                            result = st.session_state.assistant.generate_security_policy(
                                policy_org,
                                policy_focus
                            )
                            st.markdown(result)
    elif tabs == 'Malware Analysis':
        st.title("Malware Analyzer")
        tab1, tab2, tab3 = st.tabs(["📝 Paste Code", "📁 Upload Source File", "💾 Upload Binary"])

        with tab1:
            code_input = st.text_area("Paste your code here:", height=300)
            if st.button("🔍 Analyze Code", key="analyze_pasted") and code_input:
                with st.spinner("🔄 Analyzing code..."):
                    analysis_results = st.session_state.app.analyze_code(code_input)
                    display_analysis_results(analysis_results)
                    report_filename = st.session_state.app.create_analysis_report(analysis_results)
                    with open(report_filename, "rb") as file:
                        st.download_button(
                            label="📥 Download Analysis Report",
                            data=file,
                            file_name=report_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.remove(report_filename)

        with tab2:
            uploaded_file = st.file_uploader("Choose a source code file", type=['py', 'js', 'java', 'cpp', 'cs', 'php', 'rb'], key="code_file")
            if st.button("🔍 Analyze Source File", key="analyze_source") and uploaded_file:
                with st.spinner("🔄 Analyzing source file..."):
                    code_content = uploaded_file.read().decode()
                    analysis_results = st.session_state.app.analyze_code(code_content)
                    display_analysis_results(analysis_results)
                    report_filename = st.session_state.app.create_analysis_report(analysis_results)
                    with open(report_filename, "rb") as file:
                        st.download_button(
                            label="📥 Download Analysis Report",
                            data=file,
                            file_name=report_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.remove(report_filename)
        
        with tab3:
            st.write("Upload a binary file (.exe) for security analysis")
            uploaded_binary = st.file_uploader("Choose a binary file", type=['exe'], key="binary_file")
            
            if uploaded_binary:
                st.info("Binary analysis will extract strings from the executable and analyze them for security indicators.")
            
            if st.button("🔍 Analyze Binary", key="analyze_binary") and uploaded_binary:
                with st.spinner("🔄 Extracting strings and analyzing binary..."):
                    binary_data = uploaded_binary.read()
                    
                    # Create expandable section to show extracted strings
                    with st.expander("Extracted Strings Preview"):
                        extracted_strings = st.session_state.app.extract_strings_from_binary(binary_data)
                        st.text_area("Strings from binary", value=extracted_strings[:10000] + 
                                    ("\n\n[Truncated...]" if len(extracted_strings) > 10000 else ""), 
                                    height=300, disabled=True)
                    
                    analysis_results = st.session_state.app.analyze_binary(binary_data)
                    st.subheader("Binary Analysis Results")
                    display_analysis_results(analysis_results)
                    
                    report_filename = st.session_state.app.create_analysis_report(
                        analysis_results, 
                        title=f"Binary Security Analysis Report - {uploaded_binary.name}"
                    )
                    
                    with open(report_filename, "rb") as file:
                        st.download_button(
                            label="📥 Download Binary Analysis Report",
                            data=file,
                            file_name=report_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.remove(report_filename)


if __name__ == "__main__":
    main()
